import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import pandas as pd
from pathlib import Path
import fitz  # PyMuPDF

# Chemin vers Tesseract (OCR)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


class ExtracteurDocuments:

    @staticmethod
    def extract_text_pdf(pdf_path):
        text = ""
        tables = []

        # 1. Utiliser pdfplumber pour le texte et les tables
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # Tentative d'extraction du texte natif
                page_text = page.extract_text()

                # Si le PDF n'est pas texte (scanné), page_text sera None ou vide
                if page_text and page_text.strip():
                    text += page_text + "\n"
                else:
                    # Fallback : convertir la page en image avec PyMuPDF et OCR
                    # Ouvrir le PDF avec fitz
                    doc = fitz.open(pdf_path)
                    if page_num < len(doc):
                        pix = doc[page_num].get_pixmap(dpi=200)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        ocr_text = pytesseract.image_to_string(img, lang='fra')
                        text += ocr_text + "\n"
                    doc.close()

                # Extraction des tables (reste géré par pdfplumber)
                for tab in page.extract_tables():
                    if tab:
                        tables.append(pd.DataFrame(tab))

        return text, tables

    @staticmethod
    def extract_text_docx(docx_path):
        doc = Document(docx_path)
        text = "\n".join([p.text for p in doc.paragraphs])

        tables = []
        for table in doc.tables:
            data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(pd.DataFrame(data))

        return text, tables

    @staticmethod
    def extract_text_image(image_path):
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang='fra')
        return text, []

    @staticmethod
    def detect_and_extract(file_path):
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return ExtracteurDocuments.extract_text_pdf(file_path)

        elif ext in ['.docx', '.doc']:
            return ExtracteurDocuments.extract_text_docx(file_path)

        elif ext in ['.png', '.jpg', '.jpeg', '.tiff']:
            return ExtracteurDocuments.extract_text_image(file_path)

        else:
            raise ValueError(f"Format non supporté: {ext}")